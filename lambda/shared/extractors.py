"""
extractors.py - file-type routed text extraction.

Every extractor returns (raw_text, units, method) where units is a list of
    {"kind": "page"|"sheet"|"section", "index": int,
     "char_start": int, "char_end": int, "label": str|None}

`units` is the evidence substrate (EV-01). Extraction later returns a unit
index per value; that index resolves against this list to a character range.

Stage 1 scope: native formats plus digital-text PDFs. Scanned PDFs are
rejected by the readability gate rather than sent to Textract - Textract is
deferred, see README.
"""

import io
import json
import xml.dom.minidom as minidom

import docx
from docx.oxml.ns import qn
import openpyxl
from pypdf import PdfReader

# Average extractable characters per page below which a PDF is treated as
# scanned. Tuning number: set against a real corpus during Stage 1.
PDF_MIN_CHARS_PER_PAGE = 50

PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx", ".xlsm"}
STRUCTURAL_EXTS = {".json", ".xml"}


class UnreadableDocument(Exception):
    """Raised when a document has no usable text layer."""

    def __init__(self, reason):
        self.reason = reason
        super().__init__(reason)


def ext_of(key):
    tail = key.rsplit("/", 1)[-1]
    dot = tail.rfind(".")
    return tail[dot:].lower() if dot != -1 else ""


def _units_from_parts(parts, kind, labels=None):
    """Join an ordered list of per-unit text into one string plus the unit
    list. Character offsets account for the newline used to join."""
    if not parts:
        return "", [{"kind": kind, "index": 1, "char_start": 0,
                     "char_end": 0, "label": None}]

    units, cursor, joined = [], 0, []
    for i, text in enumerate(parts, start=1):
        start = cursor
        joined.append(text)
        cursor += len(text)
        units.append({
            "kind": kind,
            "index": i,
            "char_start": start,
            "char_end": cursor,
            "label": labels[i - 1] if labels else None,
        })
        cursor += 1  # the "\n" inserted by join

    return "\n".join(joined), units


def _single_unit(text, kind="section"):
    return [{"kind": kind, "index": 1, "char_start": 0,
             "char_end": len(text), "label": None}]


def extract_plain_text(body):
    text = body.decode("utf-8", errors="replace")
    return text, _single_unit(text), "plain-text"


def extract_structural(body):
    raw = body.decode("utf-8", errors="replace")
    try:
        text = json.dumps(json.loads(raw), indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        try:
            text = minidom.parseString(raw).toprettyxml(indent="  ")
        except Exception:
            text = raw
    return text, _single_unit(text), "structural-parse"


def extract_docx(body):
    """Word. Units are Word sections: a paragraph carrying <w:sectPr> closes
    the current section. python-docx exposes section properties but no
    paragraph-to-section map, so the body XML is walked directly."""
    document = docx.Document(io.BytesIO(body))
    sections, current = [], []

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            current.append("".join(t.text or "" for t in child.iter(qn("w:t"))))
            p_pr = child.find(qn("w:pPr"))
            if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
                sections.append("\n".join(current))
                current = []
        elif child.tag == qn("w:tbl"):
            rows = []
            for tr in child.iter(qn("w:tr")):
                cells = ["".join(t.text or "" for t in tc.iter(qn("w:t")))
                         for tc in tr.iter(qn("w:tc"))]
                rows.append("\t".join(cells))
            current.append("\n".join(rows))

    if current:
        sections.append("\n".join(current))
    if not sections:
        sections = [""]

    text, units = _units_from_parts(sections, "section")
    return text, units, "python-docx"


def extract_xlsx(body):
    """Spreadsheet. Units are worksheets, labelled with the sheet name so a
    citation can say 'Revenue!B14' rather than 'unit 2'."""
    wb = openpyxl.load_workbook(io.BytesIO(body), read_only=True, data_only=True)
    parts, labels = [], []

    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append("\t".join("" if c is None else str(c) for c in row))
        parts.append("# Sheet: {}\n".format(ws.title) + "\n".join(rows))
        labels.append(ws.title)

    wb.close()

    if not parts:
        parts, labels = [""], [None]

    text, units = _units_from_parts(parts, "sheet", labels)
    return text, units, "openpyxl"


def extract_pdf(body):
    """Digital-text PDF. Units are pages.

    Raises UnreadableDocument when the text layer is absent or too thin -
    the readability gate. Stage 1 has no OCR fallback, so the operator is
    told to OCR the file and resubmit."""
    try:
        reader = PdfReader(io.BytesIO(body))
        pages = [(p.extract_text() or "") for p in reader.pages]
    except Exception as exc:
        raise UnreadableDocument("pdf-parse-failed: {}".format(exc))

    if not pages:
        raise UnreadableDocument("no-pages")

    avg = sum(len(p) for p in pages) / len(pages)
    if avg < PDF_MIN_CHARS_PER_PAGE:
        raise UnreadableDocument("no_text_layer")

    text, units = _units_from_parts(pages, "page")
    return text, units, "pdf-text"


def extract(key, body):
    """Route by file type. Returns (raw_text, units, method)."""
    ext = ext_of(key)

    if ext in PDF_EXTS:
        return extract_pdf(body)
    if ext in DOCX_EXTS:
        return extract_docx(body)
    if ext in XLSX_EXTS:
        return extract_xlsx(body)
    if ext in STRUCTURAL_EXTS:
        return extract_structural(body)
    return extract_plain_text(body)
